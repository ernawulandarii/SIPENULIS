# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'question.ui'
#
# Created by: PyQt5 UI code generator 5.15.0
#
# WARNING: Any manual changes made to this file will be lost when pyuic5 is
# run again.  Do not edit this file unless you know what you are doing.


from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QTableWidgetItem
import mysql.connector
import rabin_karp as rk
import time
import statistics
import speech_recognition as sr
import docx
import datetime
from getch import pause_exit
#from openpyxl import Workbook,load_workbook


mydb=mysql.connector.connect(host="localhost",user="root",password="",database="db_skripsi")

startTime = time.time()
n = 2
#w = 2
p = 2
student = 1
question = 2
jmlquest = 2
human_rater = [35, 73, 36, 62, 42, 75, 90, 65, 86, 68, 97, 57, 15, 77, 92, 97, 93, 96, 80, 87, 88, 98, 94, 96, 98, 79,
               79, 81, 70, 74, 37, 81, 37, 77, 90, 73, 42, 96, 87, 94, 88, 63, 72]
no1 = []
excel = []
list_score = []
text = docx.Document('soal1.docx')

class Ui_QuestForm(object):
    def __init__(self, name):
        self.name = name
        self.messagebox("Congrats", "Welcome " + self.name)
        print(self.name)
        self.text = [text.paragraphs[i] for i in range (jmlquest)]
        self.doc = self.text.text
        print(self.doc)
        self.check()

    def messagebox(self, title, message):
        mess = QtWidgets.QMessageBox()
        mess.setWindowTitle(title)
        mess.setText(message)
        mess.setStandardButtons(QtWidgets.QMessageBox.Ok)
        mess.exec_()

    def check(self):
        nameid=self.name
        cur = mydb.cursor()
        query = "SELECT * from user where name=%s"
        data = cur.execute(query, [nameid])
        rows = cur.fetchall()
        print(rows)



    def startRecord(self, a):
        test = []
        r = sr.Recognizer()
        with sr.Microphone() as source:
            for i in range(jmlquest):
                a = text.paragraphs[i].text
                self.textQuestion.setText(a)
                audio = r.listen(source)
                sample_rate_hertz = 8000
                enable_automatic_punctuation = True
                r.adjust_for_ambient_noise(source)

                try:
                    test.append(r.recognize_google(audio, language="ja-JP"))
                    print("あなたの答えは:", test)
                    print("\n")

                except:
                    print('Sorry.. run again...')


            #NO CHANGE

    def setupUi(self, RecordForm):
        RecordForm.setObjectName("RecordForm")
        RecordForm.resize(933, 640)
        self.verticalLayout = QtWidgets.QVBoxLayout(RecordForm)
        self.verticalLayout.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout.setObjectName("verticalLayout")
        self.widget_2 = QtWidgets.QWidget(RecordForm)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Ignored, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.widget_2.sizePolicy().hasHeightForWidth())
        self.widget_2.setSizePolicy(sizePolicy)
        self.widget_2.setObjectName("widget_2")
        self.verticalLayout_7 = QtWidgets.QVBoxLayout(self.widget_2)
        self.verticalLayout_7.setObjectName("verticalLayout_7")
        self.horizontalLayout_4 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_4.setContentsMargins(0, 10, -1, 10)
        self.horizontalLayout_4.setSpacing(6)
        self.horizontalLayout_4.setObjectName("horizontalLayout_4")
        self.lbl_nama = QtWidgets.QLabel(self.widget_2)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.lbl_nama.sizePolicy().hasHeightForWidth())
        self.lbl_nama.setSizePolicy(sizePolicy)
        self.lbl_nama.setStyleSheet("color: rgb(231, 231, 231);\n"
"font: 15pt \"Verdana\";")
        self.lbl_nama.setObjectName("lbl_nama")
        self.horizontalLayout_4.addWidget(self.lbl_nama)
        self.nameLabel = QtWidgets.QLabel(self.widget_2)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.nameLabel.sizePolicy().hasHeightForWidth())
        self.nameLabel.setSizePolicy(sizePolicy)
        self.nameLabel.setStyleSheet("color: rgb(231, 231, 231);\n"
"font: 15pt \"Verdana\";")
        self.nameLabel.setText("")
        self.nameLabel.setObjectName("nameLabel")
        self.horizontalLayout_4.addWidget(self.nameLabel)
        self.line_4 = QtWidgets.QFrame(self.widget_2)
        self.line_4.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_4.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_4.setObjectName("line_4")
        self.horizontalLayout_4.addWidget(self.line_4)
        self.lbl_score = QtWidgets.QLabel(self.widget_2)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.lbl_score.sizePolicy().hasHeightForWidth())
        self.lbl_score.setSizePolicy(sizePolicy)
        self.lbl_score.setStyleSheet("color: rgb(231, 231, 231);\n"
"font: 15pt \"Verdana\";")
        self.lbl_score.setObjectName("lbl_score")
        self.horizontalLayout_4.addWidget(self.lbl_score)
        self.scoreLabel = QtWidgets.QLabel(self.widget_2)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.scoreLabel.sizePolicy().hasHeightForWidth())
        self.scoreLabel.setSizePolicy(sizePolicy)
        self.scoreLabel.setStyleSheet("color: rgb(231, 231, 231);\n"
"font: 15pt \"Verdana\";")
        self.scoreLabel.setText("")
        self.scoreLabel.setAlignment(QtCore.Qt.AlignCenter)
        self.scoreLabel.setObjectName("scoreLabel")
        self.horizontalLayout_4.addWidget(self.scoreLabel)
        self.verticalLayout_7.addLayout(self.horizontalLayout_4)
        self.verticalLayout_4 = QtWidgets.QVBoxLayout()
        self.verticalLayout_4.setSpacing(0)
        self.verticalLayout_4.setObjectName("verticalLayout_4")
        self.line_5 = QtWidgets.QFrame(self.widget_2)
        self.line_5.setStyleSheet("background-color:rgb(25,25, 40);")
        self.line_5.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_5.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_5.setObjectName("line_5")
        self.verticalLayout_4.addWidget(self.line_5)
        self.line_7 = QtWidgets.QFrame(self.widget_2)
        self.line_7.setAutoFillBackground(False)
        self.line_7.setStyleSheet("background-color: rgb(25,25, 40);")
        self.line_7.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_7.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_7.setObjectName("line_7")
        self.verticalLayout_4.addWidget(self.line_7)
        self.icon = QtWidgets.QLabel(self.widget_2)
        self.icon.setText("")
        self.icon.setPixmap(QtGui.QPixmap("../../SKRIPSI NEW/Images/icon.png"))
        self.icon.setAlignment(QtCore.Qt.AlignCenter)
        self.icon.setObjectName("icon")
        self.verticalLayout_4.addWidget(self.icon)
        self.line_8 = QtWidgets.QFrame(self.widget_2)
        self.line_8.setStyleSheet("background-color:rgb(25,25, 40);")
        self.line_8.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_8.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_8.setObjectName("line_8")
        self.verticalLayout_4.addWidget(self.line_8)
        self.verticalLayout_7.addLayout(self.verticalLayout_4)
        spacerItem = QtWidgets.QSpacerItem(20, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_7.addItem(spacerItem)
        self.verticalLayout_5 = QtWidgets.QVBoxLayout()
        self.verticalLayout_5.setObjectName("verticalLayout_5")
        self.label_5 = QtWidgets.QLabel(self.widget_2)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_5.sizePolicy().hasHeightForWidth())
        self.label_5.setSizePolicy(sizePolicy)
        self.label_5.setStyleSheet("color: rgb(231, 231, 231);\n"
"font: 13pt \"Verdana\";")
        self.label_5.setAlignment(QtCore.Qt.AlignCenter)
        self.label_5.setObjectName("label_5")
        self.verticalLayout_5.addWidget(self.label_5)
        self.textQuestion = QtWidgets.QTextBrowser(self.widget_2)
        self.textQuestion.setObjectName("textQuestion")
        self.verticalLayout_5.addWidget(self.textQuestion)
        self.line_6 = QtWidgets.QFrame(self.widget_2)
        self.line_6.setAutoFillBackground(False)
        self.line_6.setStyleSheet("background-color:rgb(25,25, 40);")
        self.line_6.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_6.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_6.setObjectName("line_6")
        self.verticalLayout_5.addWidget(self.line_6)
        spacerItem1 = QtWidgets.QSpacerItem(20, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_5.addItem(spacerItem1)
        self.label_7 = QtWidgets.QLabel(self.widget_2)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_7.sizePolicy().hasHeightForWidth())
        self.label_7.setSizePolicy(sizePolicy)
        self.label_7.setStyleSheet("color: rgb(231, 231, 231);\n"
"font: 12pt \"Verdana\";")
        self.label_7.setAlignment(QtCore.Qt.AlignCenter)
        self.label_7.setObjectName("label_7")
        self.verticalLayout_5.addWidget(self.label_7)
        self.textAnswer = QtWidgets.QTextBrowser(self.widget_2)
        self.textAnswer.setObjectName("textAnswer")
        self.verticalLayout_5.addWidget(self.textAnswer)
        spacerItem2 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_5.addItem(spacerItem2)
        spacerItem3 = QtWidgets.QSpacerItem(20, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_5.addItem(spacerItem3)
        self.verticalLayout_7.addLayout(self.verticalLayout_5)
        self.horizontalLayout_5 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_5.setSpacing(30)
        self.horizontalLayout_5.setObjectName("horizontalLayout_5")
        spacerItem4 = QtWidgets.QSpacerItem(350, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_5.addItem(spacerItem4)
        self.recordButton = QtWidgets.QPushButton(self.widget_2)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.recordButton.sizePolicy().hasHeightForWidth())
        self.recordButton.setSizePolicy(sizePolicy)
        self.recordButton.setStyleSheet("color: rgb(231, 231, 231);\n"
"font: 12pt \"Verdana\";\n"
"border: 2px solid rgb(255, 130, 47);;\n"
"padding: 10px;\n"
"border-radius: 3px;\n"
"opacity: 200;\n"
"")
        self.recordButton.setObjectName("recordButton")
        self.horizontalLayout_5.addWidget(self.recordButton)
        self.stopButton = QtWidgets.QPushButton(self.widget_2)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.stopButton.sizePolicy().hasHeightForWidth())
        self.stopButton.setSizePolicy(sizePolicy)
        self.stopButton.setStyleSheet("color: rgb(231, 231, 231);\n"
"font: 12pt \"Verdana\";\n"
"border: 2px solid rgb(255, 130, 47);;\n"
"padding: 10px;\n"
"border-radius: 3px;\n"
"opacity: 200;\n"
"")
        self.stopButton.setObjectName("stopButton")
        self.horizontalLayout_5.addWidget(self.stopButton)
        spacerItem5 = QtWidgets.QSpacerItem(350, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_5.addItem(spacerItem5)
        self.verticalLayout_7.addLayout(self.horizontalLayout_5)
        spacerItem6 = QtWidgets.QSpacerItem(20, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        self.verticalLayout_7.addItem(spacerItem6)
        self.horizontalLayout_8 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_8.setSpacing(30)
        self.horizontalLayout_8.setObjectName("horizontalLayout_8")
        spacerItem7 = QtWidgets.QSpacerItem(400, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_8.addItem(spacerItem7)
        self.finishButton = QtWidgets.QPushButton(self.widget_2)
        self.finishButton.setStyleSheet("color: rgb(231, 231, 231);\n"
"font: 12pt \"Verdana\";\n"
"border: 2px solid rgb(255, 130, 47);;\n"
"padding: 10px;\n"
"border-radius: 3px;\n"
"opacity: 200;\n"
"")
        self.finishButton.setObjectName("finishButton")
        self.horizontalLayout_8.addWidget(self.finishButton)
        spacerItem8 = QtWidgets.QSpacerItem(400, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_8.addItem(spacerItem8)
        self.verticalLayout_7.addLayout(self.horizontalLayout_8)
        self.logoutButton = QtWidgets.QPushButton(self.widget_2)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.logoutButton.sizePolicy().hasHeightForWidth())
        self.logoutButton.setSizePolicy(sizePolicy)
        self.logoutButton.setMinimumSize(QtCore.QSize(80, 30))
        self.logoutButton.setMaximumSize(QtCore.QSize(80, 50))
        self.logoutButton.setLayoutDirection(QtCore.Qt.RightToLeft)
        self.logoutButton.setStyleSheet("border: 2px solid  rgb(13, 58, 71);")
        self.logoutButton.setObjectName("logoutButton")
        self.verticalLayout_7.addWidget(self.logoutButton)
        self.verticalLayout.addWidget(self.widget_2)

        self.retranslateUi(RecordForm)
        QtCore.QMetaObject.connectSlotsByName(RecordForm)
        RecordForm.setStyleSheet(
            """
            QWidget{background-color: rgb(25,25, 40);}
            QPushButton {
                border-style: outset;
                border-radius: 0px;
                padding: 6px;
            }
            QPushButton:hover {
                background-color: #cf7500;
                border-style: inset;
            }
            QPushButton:pressed {
                background-color: #ffa126;
                border-style: inset;
            }
            """
        )
        self.recordButton.clicked.connect(self.startRecord)

    def retranslateUi(self, RecordForm):
        _translate = QtCore.QCoreApplication.translate
        RecordForm.setWindowTitle(_translate("RecordForm", "Form"))
        self.lbl_nama.setText(_translate("RecordForm", "Name :"))
        self.lbl_score.setText(_translate("RecordForm", "Score :"))
        self.label_5.setText(_translate("RecordForm", "Answer The Questions"))
        self.label_7.setText(_translate("RecordForm", "Your Answer of Number "))
        self.recordButton.setText(_translate("RecordForm", "Record"))
        self.stopButton.setText(_translate("RecordForm", "Stop"))
        self.finishButton.setText(_translate("RecordForm", "Finish"))
        self.logoutButton.setText(_translate("RecordForm", "LOGOUT"))
        ########################################################
        self.nameLabel.setText(_translate("QuestForm", self.name))
        self.scoreLabel.setText(_translate("QuestForm", ""))


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    QuestForm = QtWidgets.QWidget()
    ui = Ui_QuestForm()
    ui.setupUi(QuestForm)
    QuestForm.show()
    sys.exit(app.exec_())
